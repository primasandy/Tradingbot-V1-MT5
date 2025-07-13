# -*- coding: utf-8 -*-

"""
Aplikasi Transkripsi Audio ke Teks Powerful dengan Python
=========================================================
Versi 4.0 - Antarmuka dengan Mode Ganda (Tunggal & Batch)
Deskripsi:
Aplikasi ini kini memiliki dua mode dalam satu jendela:
1. Transkripsi Tunggal: Untuk memproses satu file dengan fitur detail.
2. Proses Batch: Untuk otomatisasi antrian banyak file.

Dependensi:
- PyQt6: pip install PyQt6
- openai-whisper: pip install openai-whisper
- python-docx: pip install python-docx
- ffmpeg: Harus terinstal di sistem dan dapat diakses melalui PATH.

Cara Menjalankan:
Simpan kode ini sebagai file .py dan jalankan dari terminal: python nama_file.py
"""

import sys
import os
import whisper
import shutil
import base64
import docx
from PyQt6.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QPushButton, QLabel, QFileDialog,
    QTextEdit, QProgressBar, QComboBox, QMessageBox, QHBoxLayout, QGroupBox,
    QSizePolicy, QCheckBox, QListWidget, QListWidgetItem, QTabWidget
)
from PyQt6.QtCore import QThread, pyqtSignal, Qt
from PyQt6.QtGui import QFont, QIcon, QPixmap

# --- Data Ikon SVG (Base64 Encoded) ---
APP_ICON_SVG = b'PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIyNCIgaGVpZ2h0PSIyNCIgdmlld0JveD0iMCAwIDI0IDI0IiBmaWxsPSJub25lIiBzdHJva2U9IiNmZmZmZmYiIHN0cm9rZS13aWR0aD0iMiIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIiBzdHJva2UtbGluZWpvaW49InJvdW5kIiBjbGFzcz0ibHVjaWRlIGx1Y2lkZS1hdWRpby13YXZlZm9ybSI+PHBhdGggZD0iTTMgMTJoMlY4aC0yem00IDBoMlY0aC0yek0xMSAxMmgyVjBoLTJ6bTQgMGgyVjZoLTJ6bTE5IDEyaDJWOGgtMnoiLz48L3N2Zz4='

# --- Worker Thread untuk Proses Transkripsi (Bisa dipakai kedua mode) ---
class Worker(QThread):
    finished = pyqtSignal(dict) # Sinyal untuk mode tunggal
    finished_file = pyqtSignal(str, dict) # Sinyal untuk mode batch
    progress = pyqtSignal(int)
    error = pyqtSignal(str)

    def __init__(self, file_path, model_name, language):
        super().__init__()
        self.file_path = file_path
        self.model_name = model_name
        self.language = language

    def run(self):
        try:
            self.progress.emit(10)
            model = whisper.load_model(self.model_name)
            self.progress.emit(30)
            options = {"fp16": False}
            if self.language: options["language"] = self.language
            self.progress.emit(50)
            result = model.transcribe(self.file_path, **options)
            self.progress.emit(90)
            self.finished.emit(result) # Kirim sinyal untuk mode tunggal
            self.finished_file.emit(self.file_path, result) # Kirim sinyal untuk mode batch
        except Exception as e:
            self.error.emit(f"Terjadi kesalahan pada file {os.path.basename(self.file_path)}: {str(e)}")

# --- TAB 1: MODE TRANSKRIPSI TUNGGAL ---
class SingleFileTab(QWidget):
    def __init__(self):
        super().__init__()
        self.file_path = None
        self.transcription_result = None
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(15, 15, 15, 15)

        # Kontrol
        control_group = QGroupBox("Pengaturan")
        control_layout = QVBoxLayout(control_group)
        
        file_layout = QHBoxLayout()
        self.select_button = QPushButton('📂 Pilih File Audio')
        self.select_button.clicked.connect(self.select_file)
        self.file_label = QLabel('Belum ada file yang dipilih.')
        file_layout.addWidget(self.select_button)
        file_layout.addWidget(self.file_label, 1)
        control_layout.addLayout(file_layout)

        options_layout = QHBoxLayout()
        lang_label = QLabel("Bahasa:")
        self.lang_combo = QComboBox()
        supported_langs = sorted(whisper.tokenizer.LANGUAGES.items(), key=lambda x: x[1])
        self.lang_combo.addItem("Auto-Detect", None)
        for code, name in supported_langs: self.lang_combo.addItem(f"{name.title()} ({code})", code)
        model_label = QLabel("Model:")
        self.model_combo = QComboBox()
        self.model_combo.addItems(['tiny', 'base', 'small', 'medium', 'large'])
        options_layout.addWidget(lang_label); options_layout.addWidget(self.lang_combo, 1)
        options_layout.addSpacing(20); options_layout.addWidget(model_label)
        options_layout.addWidget(self.model_combo, 1)
        control_layout.addLayout(options_layout)
        layout.addWidget(control_group)

        # Tombol Transkripsi
        self.transcribe_button = QPushButton('🎙️ Mulai Transkripsi')
        self.transcribe_button.setObjectName("transcribeButtonSingle")
        self.transcribe_button.setFixedHeight(45)
        self.transcribe_button.clicked.connect(self.start_transcription)
        layout.addWidget(self.transcribe_button)

        # Progress Bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)

        # Hasil
        result_group = QGroupBox("Hasil Transkripsi")
        result_layout = QVBoxLayout(result_group)
        display_options_layout = QHBoxLayout()
        self.timestamp_checkbox = QCheckBox("Tampilkan Stempel Waktu")
        self.timestamp_checkbox.stateChanged.connect(self.update_display)
        self.word_count_label = QLabel("Kata: 0 | Karakter: 0")
        self.word_count_label.setAlignment(Qt.AlignmentFlag.AlignRight)
        display_options_layout.addWidget(self.timestamp_checkbox)
        display_options_layout.addStretch()
        display_options_layout.addWidget(self.word_count_label)
        result_layout.addLayout(display_options_layout)

        self.result_text = QTextEdit()
        self.result_text.setReadOnly(True)
        result_layout.addWidget(self.result_text)

        action_buttons_layout = QHBoxLayout()
        action_buttons_layout.addStretch()
        self.copy_button = QPushButton("Salin Teks")
        self.copy_button.clicked.connect(self.copy_to_clipboard)
        self.export_button = QPushButton("Ekspor Hasil")
        self.export_button.clicked.connect(self.export_to_file)
        action_buttons_layout.addWidget(self.copy_button)
        action_buttons_layout.addWidget(self.export_button)
        result_layout.addLayout(action_buttons_layout)
        layout.addWidget(result_group)

    def select_file(self):
        file, _ = QFileDialog.getOpenFileName(self, "Pilih File Audio", "", "Audio Files (*.mp3 *.wav *.m4a *.flac *.ogg *.opus)")
        if file:
            self.file_path = file
            self.file_label.setText(f"File: {os.path.basename(file)}")
            self.result_text.clear()
            self.transcription_result = None
            self.update_display()

    def start_transcription(self):
        if not self.check_ffmpeg() or not self.file_path: return
        self.transcribe_button.setEnabled(False)
        self.progress_bar.setVisible(True)
        self.worker = Worker(self.file_path, self.model_combo.currentText(), self.lang_combo.currentData())
        self.worker.progress.connect(self.progress_bar.setValue)
        self.worker.finished.connect(self.on_transcription_finished)
        self.worker.error.connect(self.on_transcription_error)
        self.worker.start()

    def on_transcription_finished(self, result):
        self.transcription_result = result
        self.update_display()
        self.transcribe_button.setEnabled(True)
        self.progress_bar.setVisible(False)
        QMessageBox.information(self, 'Selesai', 'Transkripsi berhasil diselesaikan.')

    def on_transcription_error(self, error_message):
        QMessageBox.critical(self, 'Error', error_message)
        self.transcribe_button.setEnabled(True)
        self.progress_bar.setVisible(False)

    def update_display(self):
        if not self.transcription_result:
            self.result_text.clear()
            self.word_count_label.setText("Kata: 0 | Karakter: 0")
            return
        
        text_to_display = ""
        if self.timestamp_checkbox.isChecked():
            for segment in self.transcription_result['segments']:
                start, text = int(segment['start']), segment['text'].strip()
                text_to_display += f"[{start//60:02d}:{start%60:02d}] {text}\n"
        else:
            text_to_display = self.transcription_result['text']
        
        self.result_text.setText(text_to_display)
        self.word_count_label.setText(f"Kata: {len(text_to_display.split())} | Karakter: {len(text_to_display)}")

    def copy_to_clipboard(self):
        QApplication.clipboard().setText(self.result_text.toPlainText())

    def export_to_file(self):
        if not self.result_text.toPlainText(): return
        path, flt = QFileDialog.getSaveFileName(self, "Simpan Hasil", "", "Dokumen Word (*.docx);;File Teks (*.txt)")
        if path:
            try:
                if flt.startswith("Dokumen Word"):
                    doc = docx.Document()
                    doc.add_paragraph(self.result_text.toPlainText())
                    doc.save(path)
                else:
                    with open(path, 'w', encoding='utf-8') as f: f.write(self.result_text.toPlainText())
                QMessageBox.information(self, "Sukses", "File berhasil disimpan!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Gagal menyimpan file:\n{e}")

    def check_ffmpeg(self):
        if shutil.which("ffmpeg"): return True
        QMessageBox.critical(self, "Error: FFmpeg Tidak Ditemukan", "FFmpeg tidak ditemukan.")
        return False

# --- TAB 2: MODE PROSES BATCH ---
class BatchProcessingTab(QWidget):
    def __init__(self):
        super().__init__()
        self.is_processing = False
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(15, 15, 15, 15)

        # Grup Kontrol Utama
        group_box = QGroupBox("Pengaturan & Antrian File")
        main_layout = QHBoxLayout(group_box)
        
        # Kolom Kiri: Antrian
        queue_layout = QVBoxLayout()
        self.file_list_widget = QListWidget()
        queue_layout.addWidget(QLabel("Daftar Antrian Audio:"))
        queue_layout.addWidget(self.file_list_widget)
        queue_buttons_layout = QHBoxLayout()
        add_file_btn = QPushButton("Tambah File"); add_folder_btn = QPushButton("Tambah Folder")
        remove_btn = QPushButton("Hapus"); clear_btn = QPushButton("Kosongkan")
        add_file_btn.clicked.connect(self.add_files); add_folder_btn.clicked.connect(self.add_folder)
        remove_btn.clicked.connect(self.remove_selected_file); clear_btn.clicked.connect(self.clear_queue)
        queue_buttons_layout.addWidget(add_file_btn); queue_buttons_layout.addWidget(add_folder_btn)
        queue_buttons_layout.addStretch(); queue_buttons_layout.addWidget(remove_btn); queue_buttons_layout.addWidget(clear_btn)
        queue_layout.addLayout(queue_buttons_layout)
        
        # Kolom Kanan: Opsi
        options_layout = QVBoxLayout()
        lang_model_layout = QHBoxLayout()
        self.lang_combo = QComboBox()
        supported_langs = sorted(whisper.tokenizer.LANGUAGES.items(), key=lambda x: x[1])
        self.lang_combo.addItem("Auto-Detect", None)
        for code, name in supported_langs: self.lang_combo.addItem(f"{name.title()} ({code})", code)
        self.model_combo = QComboBox()
        self.model_combo.addItems(['tiny', 'base', 'small', 'medium', 'large'])
        lang_model_layout.addWidget(QLabel("Bahasa:")); lang_model_layout.addWidget(self.lang_combo)
        lang_model_layout.addWidget(QLabel("Model:")); lang_model_layout.addWidget(self.model_combo)
        options_layout.addLayout(lang_model_layout)
        
        export_group = QGroupBox("Format Ekspor Otomatis")
        export_layout = QHBoxLayout(export_group)
        self.export_txt_check = QCheckBox(".txt"); self.export_txt_check.setChecked(True)
        self.export_docx_check = QCheckBox(".docx"); self.export_srt_check = QCheckBox(".srt")
        export_layout.addWidget(self.export_txt_check); export_layout.addWidget(self.export_docx_check); export_layout.addWidget(self.export_srt_check)
        options_layout.addWidget(export_group)
        options_layout.addStretch()
        
        main_layout.addLayout(queue_layout, 2); main_layout.addLayout(options_layout, 1)
        layout.addWidget(group_box)

        # Tombol Proses
        self.transcribe_button = QPushButton('🚀 Mulai Proses Antrian')
        self.transcribe_button.setObjectName("transcribeButtonBatch")
        self.transcribe_button.setFixedHeight(45)
        self.transcribe_button.clicked.connect(self.start_batch_processing)
        layout.addWidget(self.transcribe_button)

        # Log
        self.result_log = QTextEdit()
        self.result_log.setReadOnly(True)
        layout.addWidget(QGroupBox("Log Proses", objectName="logGroup"))
        layout.itemAt(layout.count()-1).widget().setLayout(QVBoxLayout())
        layout.itemAt(layout.count()-1).widget().layout().addWidget(self.result_log)

    def add_files(self):
        files, _ = QFileDialog.getOpenFileNames(self, "Pilih File Audio", "", "Audio Files (*.mp3 *.wav *.m4a *.flac *.ogg *.opus)")
        for file in files:
            if not self.file_list_widget.findItems(file, Qt.MatchFlag.MatchExactly):
                self.file_list_widget.addItem(file)

    def add_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Pilih Folder")
        if folder:
            for root, _, files in os.walk(folder):
                for file in files:
                    if file.lower().endswith(('.mp3', '.wav', '.m4a', '.flac')):
                        path = os.path.join(root, file)
                        if not self.file_list_widget.findItems(path, Qt.MatchFlag.MatchExactly):
                            self.file_list_widget.addItem(path)

    def remove_selected_file(self):
        for item in self.file_list_widget.selectedItems(): self.file_list_widget.takeItem(self.file_list_widget.row(item))
    def clear_queue(self): self.file_list_widget.clear()

    def start_batch_processing(self):
        if self.is_processing or self.file_list_widget.count() == 0 or not self.check_ffmpeg(): return
        self.is_processing = True
        self.file_queue = [self.file_list_widget.item(i).text() for i in range(self.file_list_widget.count())]
        self.current_file_index = 0
        self.result_log.clear()
        self.transcribe_button.setText("Menghentikan Proses..."); self.transcribe_button.clicked.disconnect(); self.transcribe_button.clicked.connect(self.stop_processing)
        self.process_next_file()

    def stop_processing(self):
        self.is_processing = False
        self.transcribe_button.setText("🚀 Mulai Proses Antrian"); self.transcribe_button.clicked.disconnect(); self.transcribe_button.clicked.connect(self.start_batch_processing)

    def process_next_file(self):
        if not self.is_processing or self.current_file_index >= len(self.file_queue):
            self.stop_processing()
            QMessageBox.information(self, "Selesai", "Proses batch selesai!")
            return
        file_path = self.file_queue[self.current_file_index]
        self.result_log.append(f"▶️ Memproses {os.path.basename(file_path)}...")
        self.worker = Worker(file_path, self.model_combo.currentText(), self.lang_combo.currentData())
        self.worker.finished_file.connect(self.on_file_finished)
        self.worker.error.connect(self.on_file_error)
        self.worker.start()

    def on_file_finished(self, file_path, result):
        if not self.is_processing: return
        self.result_log.append(f"✅ Selesai: {os.path.basename(file_path)}")
        self.save_results(file_path, result)
        self.current_file_index += 1
        self.process_next_file()

    def on_file_error(self, error_message):
        if not self.is_processing: return
        self.result_log.append(f"❌ GAGAL: {error_message}")
        self.current_file_index += 1
        self.process_next_file()

    def save_results(self, original_path, result):
        output_dir = "Hasil Transkripsi"
        os.makedirs(output_dir, exist_ok=True)
        base_filename = os.path.splitext(os.path.basename(original_path))[0]
        if self.export_txt_check.isChecked():
            with open(os.path.join(output_dir, f"{base_filename}.txt"), 'w', encoding='utf-8') as f: f.write(result['text'])
        if self.export_docx_check.isChecked():
            doc = docx.Document(); doc.add_paragraph(result['text']); doc.save(os.path.join(output_dir, f"{base_filename}.docx"))
        if self.export_srt_check.isChecked():
            with open(os.path.join(output_dir, f"{base_filename}.srt"), 'w', encoding='utf-8') as f:
                for i, seg in enumerate(result['segments']):
                    f.write(f"{i+1}\n{self.format_srt_time(seg['start'])} --> {self.format_srt_time(seg['end'])}\n{seg['text'].strip()}\n\n")
        self.result_log.append(f"   ↳ Hasil disimpan di folder '{output_dir}'")

    def format_srt_time(self, s):
        h, m = divmod(s, 3600); m, s = divmod(m, 60); ms = int((s - int(s)) * 1000)
        return f"{int(h):02d}:{int(m):02d}:{int(s):02d},{ms:03d}"

    def check_ffmpeg(self):
        if shutil.which("ffmpeg"): return True
        QMessageBox.critical(self, "Error: FFmpeg Tidak Ditemukan", "FFmpeg tidak ditemukan.")
        return False

# --- Jendela Utama Aplikasi dengan Tabs ---
class TranscriberApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('Audio Scribe Pro - Otomatisasi Transkripsi')
        self.setGeometry(100, 100, 850, 750)
        
        pixmap = QPixmap(); pixmap.loadFromData(base64.b64decode(APP_ICON_SVG))
        self.setWindowIcon(QIcon(pixmap))

        main_layout = QVBoxLayout(self)
        self.tabs = QTabWidget()
        self.tabs.addTab(SingleFileTab(), "Transkripsi Tunggal")
        self.tabs.addTab(BatchProcessingTab(), "Proses Batch")
        main_layout.addWidget(self.tabs)
        self.apply_styles()

    def apply_styles(self):
        self.setStyleSheet("""
            QWidget { background-color: #2c3e50; color: #ecf0f1; font-family: 'Segoe UI', Arial, sans-serif; }
            QTabWidget::pane { border: 1px solid #34495e; }
            QTabBar::tab {
                background: #34495e; color: #ecf0f1; padding: 12px 25px;
                border: 1px solid #34495e; border-bottom: none;
                border-top-left-radius: 6px; border-top-right-radius: 6px;
                font-weight: bold;
            }
            QTabBar::tab:selected { background: #2c3e50; border-color: #3498db; }
            QTabBar::tab:!selected:hover { background: #4a627a; }
            QGroupBox { font-size: 14px; font-weight: bold; border: 1px solid #34495e; border-radius: 8px; margin-top: 10px; padding: 15px; }
            QGroupBox::title { subcontrol-origin: margin; subcontrol-position: top left; padding: 0 10px; }
            QLabel { font-size: 12px; }
            QPushButton { background-color: #3498db; color: white; font-size: 12px; font-weight: bold; border: none; border-radius: 5px; padding: 10px 15px; }
            QPushButton:hover { background-color: #2980b9; }
            QPushButton:disabled { background-color: #566573; }
            #transcribeButtonSingle { background-color: #27ae60; font-size: 16px; padding: 12px; }
            #transcribeButtonSingle:hover { background-color: #229954; }
            #transcribeButtonBatch { background-color: #e67e22; font-size: 16px; padding: 12px; }
            #transcribeButtonBatch:hover { background-color: #d35400; }
            QComboBox, QListWidget, QTextEdit { border: 1px solid #34495e; border-radius: 5px; padding: 8px; background-color: #34495e; }
            QProgressBar { border: none; border-radius: 4px; background-color: #34495e; text-align: center; color: white; }
            QProgressBar::chunk { background-color: #27ae60; border-radius: 4px; }
            QCheckBox { spacing: 5px; }
        """)

if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = TranscriberApp()
    ex.show()
    sys.exit(app.exec())
